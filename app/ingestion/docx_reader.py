from dataclasses import dataclass, field
from pathlib import Path

from docx import Document

from app.ingestion.file_validator import validate_docx_file


@dataclass(frozen=True)
class ExtractedTable:
    rows: list[list[str]]


@dataclass(frozen=True)
class ExtractedDocxContent:
    source_path: Path
    paragraphs: list[str]
    tables: list[ExtractedTable] = field(default_factory=list)

    @property
    def plain_text(self) -> str:
        parts: list[str] = []
        parts.extend(self.paragraphs)
        for table in self.tables:
            for row in table.rows:
                parts.append(" | ".join(cell for cell in row if cell))
        return "\n".join(part for part in parts if part).strip()


def read_docx(path: str | Path) -> ExtractedDocxContent:
    file_path = validate_docx_file(path)
    document = Document(str(file_path))

    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    tables: list[ExtractedTable] = []

    for table in document.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(cells)
        if rows:
            tables.append(ExtractedTable(rows=rows))

    return ExtractedDocxContent(source_path=file_path, paragraphs=paragraphs, tables=tables)
