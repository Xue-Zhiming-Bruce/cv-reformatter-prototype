from pathlib import Path
from typing import Mapping

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph

from app.generation.target_format_blueprint import (
    BlueprintSection,
    TargetFormatBlueprint,
    TextStyle,
    get_target_format_blueprint,
)
from app.generation.template_mapper import ClientFacingRenderContext, RenderEducation, RenderWorkExperience


DEFAULT_OUTPUT_DIR = Path("data/generated_outputs")


class DocxRenderingError(RuntimeError):
    """Raised when a client-facing DOCX cannot be rendered."""


def render_docx(
    context: ClientFacingRenderContext | Mapping[str, object],
    output_path: str | Path | None = None,
    *,
    output_dir: str | Path = DEFAULT_OUTPUT_DIR,
    blueprint: TargetFormatBlueprint | None = None,
) -> Path:
    render_context = _coerce_context(context)
    resolved_blueprint = blueprint or get_target_format_blueprint(render_context.template_name)
    resolved_output_path = _resolve_output_path(output_path, output_dir)
    resolved_output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    _configure_document_from_blueprint(document, resolved_blueprint)
    if resolved_blueprint.renderer_key == "classic_10554236":
        _render_classic_10554236(document, render_context, resolved_blueprint)
    else:
        _render_apex_standard(document, render_context)

    try:
        document.save(resolved_output_path)
    except Exception as exc:  # pragma: no cover - python-docx raises library-specific exceptions.
        raise DocxRenderingError(f"DOCX could not be written to {resolved_output_path}") from exc
    return resolved_output_path


def _coerce_context(context: ClientFacingRenderContext | Mapping[str, object]) -> ClientFacingRenderContext:
    if isinstance(context, ClientFacingRenderContext):
        return context
    return ClientFacingRenderContext.model_validate(context)


def _resolve_output_path(output_path: str | Path | None, output_dir: str | Path) -> Path:
    if output_path is not None:
        return Path(output_path)
    return Path(output_dir) / "candidate_profile.docx"


def _configure_document_from_blueprint(document: DocxDocument, blueprint: TargetFormatBlueprint) -> None:
    section = document.sections[0]
    section.page_width = Inches(blueprint.page.width_inches)
    section.page_height = Inches(blueprint.page.height_inches)
    section.top_margin = Inches(blueprint.page.top_margin_inches)
    section.right_margin = Inches(blueprint.page.right_margin_inches)
    section.bottom_margin = Inches(blueprint.page.bottom_margin_inches)
    section.left_margin = Inches(blueprint.page.left_margin_inches)
    section.header_distance = Inches(blueprint.page.header_distance_inches)
    section.footer_distance = Inches(blueprint.page.footer_distance_inches)

    normal = blueprint.typography["normal"]
    normal_style = document.styles["Normal"]
    normal_style.font.name = normal.font_name
    normal_style.font.size = Pt(normal.size_pt)
    normal_style.font.bold = normal.bold
    normal_style.font.italic = normal.italic
    normal_style.font.color.rgb = RGBColor.from_string(normal.color_hex)
    normal_style.paragraph_format.line_spacing = normal.line_spacing
    normal_style.paragraph_format.space_before = Pt(normal.space_before_pt)
    normal_style.paragraph_format.space_after = Pt(normal.space_after_pt)
    normal_style.paragraph_format.widow_control = blueprint.overflow.widow_control

    for style_name, role in [("Title", "title"), ("Heading 1", "section"), ("Heading 2", "section")]:
        style_definition = blueprint.typography[role]
        style = document.styles[style_name]
        style.font.name = style_definition.font_name
        style.font.size = Pt(style_definition.size_pt)
        style.font.bold = style_definition.bold
        style.font.italic = style_definition.italic
        style.font.color.rgb = RGBColor.from_string(style_definition.color_hex)
        style.paragraph_format.line_spacing = style_definition.line_spacing
        style.paragraph_format.space_before = Pt(style_definition.space_before_pt)
        style.paragraph_format.space_after = Pt(style_definition.space_after_pt)
        style.paragraph_format.keep_with_next = blueprint.overflow.keep_section_heading_with_next


def _render_apex_standard(document: DocxDocument, context: ClientFacingRenderContext) -> None:
    _render_header(document, context)
    _render_contact_block(document, context)
    _render_summary(document, context)
    _render_simple_list_section(document, "Skills", context.skills)
    _render_simple_list_section(document, "Languages", context.languages)
    _render_experience(document, context.work_experience)
    _render_education(document, context.education)
    _render_simple_list_section(document, "Certifications", context.certifications)
    _render_additional_details(document, context)


def _render_classic_10554236(
    document: DocxDocument,
    context: ClientFacingRenderContext,
    blueprint: TargetFormatBlueprint,
) -> None:
    renderers = {
        BlueprintSection.HEADER: lambda: _classic_header(document, context, blueprint),
        BlueprintSection.CONTACT: lambda: _classic_contact(document, context, blueprint),
        BlueprintSection.SUMMARY: lambda: _classic_summary(document, context, blueprint),
        BlueprintSection.HIGHLIGHTS: lambda: _classic_highlights(document, context, blueprint),
        BlueprintSection.LANGUAGES: lambda: _classic_simple_line_section(
            document, "Languages", context.languages, blueprint
        ),
        BlueprintSection.EXPERIENCE: lambda: _classic_experience(document, context.work_experience, blueprint),
        BlueprintSection.EDUCATION: lambda: _classic_education(document, context.education, blueprint),
        BlueprintSection.CERTIFICATIONS: lambda: _classic_simple_line_section(
            document, "Certifications", context.certifications, blueprint
        ),
        BlueprintSection.ADDITIONAL_DETAILS: lambda: _classic_additional_details(document, context, blueprint),
        BlueprintSection.SKILLS: lambda: _classic_skills_line(document, context.skills, blueprint),
    }
    for section in blueprint.section_order:
        renderers[section]()


def _classic_header(
    document: DocxDocument,
    context: ClientFacingRenderContext,
    blueprint: TargetFormatBlueprint,
) -> None:
    title = context.candidate_subheading or context.candidate_heading
    if not title:
        return
    title = title.split(" | ", maxsplit=1)[0]
    paragraph = document.add_paragraph()
    paragraph.add_run(title.upper())
    _apply_paragraph_style(paragraph, blueprint.typography["title"])


def _classic_contact(
    document: DocxDocument,
    context: ClientFacingRenderContext,
    blueprint: TargetFormatBlueprint,
) -> None:
    if context.contact_lines:
        _classic_simple_line_section(document, "Contact", context.contact_lines, blueprint)


def _classic_summary(
    document: DocxDocument,
    context: ClientFacingRenderContext,
    blueprint: TargetFormatBlueprint,
) -> None:
    if not context.professional_summary:
        return
    _classic_section_heading(document, "Summary", blueprint)
    paragraph = document.add_paragraph(context.professional_summary)
    _apply_paragraph_style(paragraph, blueprint.typography["normal"])


def _classic_highlights(
    document: DocxDocument,
    context: ClientFacingRenderContext,
    blueprint: TargetFormatBlueprint,
) -> None:
    if not context.skills:
        return
    _classic_section_heading(document, "Highlights", blueprint)
    columns = min(blueprint.skills_columns, max(1, len(context.skills)))
    rows = (len(context.skills) + columns - 1) // columns
    table = document.add_table(rows=rows, cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    for index, item in enumerate(context.skills):
        row = index % rows
        column = index // rows
        cell = table.cell(row, column)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        paragraph = cell.paragraphs[0]
        paragraph.style = document.styles["List Bullet"]
        paragraph.add_run(item)
        _configure_bullet_paragraph(paragraph, blueprint)


def _classic_experience(
    document: DocxDocument,
    entries: list[RenderWorkExperience],
    blueprint: TargetFormatBlueprint,
) -> None:
    if not entries:
        return
    _classic_section_heading(document, "Experience", blueprint)
    for entry in entries:
        date_range = entry.date_range.replace(" - ", " to ") if entry.date_range else None
        heading = _join_present([entry.company, date_range, entry.title], separator=" ")
        if heading:
            paragraph = document.add_paragraph(heading)
            _apply_paragraph_style(paragraph, blueprint.typography["entry_heading"])
            paragraph.paragraph_format.keep_with_next = bool(entry.location or entry.description)
        if entry.location:
            paragraph = document.add_paragraph(entry.location)
            _apply_paragraph_style(paragraph, blueprint.typography["normal"])
            paragraph.paragraph_format.keep_with_next = bool(entry.description)
        for index, description_item in enumerate(entry.description):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(description_item)
            _configure_bullet_paragraph(paragraph, blueprint)
            if index == 0 and blueprint.overflow.keep_experience_heading_with_first_bullet:
                paragraph.paragraph_format.keep_with_next = False


def _classic_education(
    document: DocxDocument,
    entries: list[RenderEducation],
    blueprint: TargetFormatBlueprint,
) -> None:
    if not entries:
        return
    _classic_section_heading(document, "Education", blueprint)
    for entry in entries:
        date_range = entry.date_range.replace(" - ", " to ") if entry.date_range else None
        text = _join_present([entry.institution, date_range, entry.degree, entry.field_of_study], separator=" ")
        if text:
            paragraph = document.add_paragraph(text)
            _apply_paragraph_style(paragraph, blueprint.typography["normal"])


def _classic_simple_line_section(
    document: DocxDocument,
    title: str,
    items: list[str],
    blueprint: TargetFormatBlueprint,
) -> None:
    if not items:
        return
    _classic_section_heading(document, title, blueprint)
    for item in items:
        paragraph = document.add_paragraph(item)
        _apply_paragraph_style(paragraph, blueprint.typography["normal"])


def _classic_additional_details(
    document: DocxDocument,
    context: ClientFacingRenderContext,
    blueprint: TargetFormatBlueprint,
) -> None:
    if not context.additional_details:
        return
    _classic_section_heading(document, "Additional Information", blueprint)
    for detail in context.additional_details:
        paragraph = document.add_paragraph(f"{detail.label}: {detail.value}")
        _apply_paragraph_style(paragraph, blueprint.typography["normal"])


def _classic_skills_line(
    document: DocxDocument,
    skills: list[str],
    blueprint: TargetFormatBlueprint,
) -> None:
    if not skills:
        return
    _classic_section_heading(document, "Skills", blueprint)
    paragraph = document.add_paragraph("; ".join(skills) + ".")
    _apply_paragraph_style(paragraph, blueprint.typography["normal"])


def _classic_section_heading(
    document: DocxDocument,
    title: str,
    blueprint: TargetFormatBlueprint,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run(title)
    _apply_paragraph_style(paragraph, blueprint.typography["section"])
    paragraph.paragraph_format.keep_with_next = blueprint.overflow.keep_section_heading_with_next


def _configure_bullet_paragraph(paragraph: Paragraph, blueprint: TargetFormatBlueprint) -> None:
    paragraph.paragraph_format.left_indent = Inches(blueprint.bullet_left_indent_inches)
    paragraph.paragraph_format.first_line_indent = Inches(-blueprint.bullet_hanging_indent_inches)
    _apply_paragraph_style(paragraph, blueprint.typography["normal"])


def _apply_paragraph_style(paragraph: Paragraph, style: TextStyle) -> None:
    paragraph.paragraph_format.line_spacing = style.line_spacing
    paragraph.paragraph_format.space_before = Pt(style.space_before_pt)
    paragraph.paragraph_format.space_after = Pt(style.space_after_pt)
    for run in paragraph.runs:
        run.font.name = style.font_name
        run.font.size = Pt(style.size_pt)
        run.font.bold = style.bold
        run.font.italic = style.italic
        run.font.color.rgb = RGBColor.from_string(style.color_hex)


def _render_header(document: DocxDocument, context: ClientFacingRenderContext) -> None:
    document.add_heading("Candidate Profile", level=0)
    heading = document.add_paragraph()
    heading.style = document.styles["Heading 1"]
    heading.add_run(context.candidate_heading).bold = True
    if context.candidate_subheading:
        document.add_paragraph(context.candidate_subheading)


def _render_contact_block(document: DocxDocument, context: ClientFacingRenderContext) -> None:
    if not context.contact_lines:
        return
    _add_section_heading(document, "Contact")
    for line in context.contact_lines:
        document.add_paragraph(line)


def _render_summary(document: DocxDocument, context: ClientFacingRenderContext) -> None:
    if not context.professional_summary:
        return
    _add_section_heading(document, "Professional Summary")
    document.add_paragraph(context.professional_summary)


def _render_simple_list_section(document: DocxDocument, title: str, items: list[str]) -> None:
    if not items:
        return
    _add_section_heading(document, title)
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def _render_experience(document: DocxDocument, entries: list[RenderWorkExperience]) -> None:
    if not entries:
        return
    _add_section_heading(document, "Work Experience")
    for entry in entries:
        title = _join_present([entry.title, entry.company, entry.date_range])
        if title:
            paragraph = document.add_paragraph()
            paragraph.add_run(title).bold = True
        if entry.location:
            document.add_paragraph(entry.location)
        for description_item in entry.description:
            document.add_paragraph(description_item, style="List Bullet")


def _render_education(document: DocxDocument, entries: list[RenderEducation]) -> None:
    if not entries:
        return
    _add_section_heading(document, "Education")
    for entry in entries:
        text = _join_present([entry.institution, entry.degree, entry.field_of_study, entry.date_range])
        if text:
            document.add_paragraph(text)


def _render_additional_details(document: DocxDocument, context: ClientFacingRenderContext) -> None:
    if not context.additional_details:
        return
    _add_section_heading(document, "Additional Details")
    for detail in context.additional_details:
        document.add_paragraph(f"{detail.label}: {detail.value}")


def _add_section_heading(document: DocxDocument, title: str) -> None:
    document.add_heading(title, level=2)


def _join_present(values: list[str | None], separator: str = " | ") -> str | None:
    present = [value for value in values if value]
    return separator.join(present) if present else None
