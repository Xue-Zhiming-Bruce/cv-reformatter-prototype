import argparse
import json
from pathlib import Path

from docx import Document
try:
    from dotenv import load_dotenv
except ModuleNotFoundError:
    def load_dotenv() -> bool:
        return False

from app.ingestion.docx_reader import read_docx
from app.ingestion.pdf_reader import read_pdf_text
from app.extraction.llm_extractor import MockLLMClient, build_llm_client, extract_candidate_profile
from app.validation.followup_message_generator import generate_followup_message


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_INPUT = PROJECT_ROOT / "data" / "input_samples" / "synthetic_resume.docx"
OUTPUT_DIR = PROJECT_ROOT / "data" / "generated_outputs"


def main() -> None:
    load_dotenv()
    parser = argparse.ArgumentParser(description="Run the resume extraction milestone demo.")
    parser.add_argument("--input", type=Path, default=DEFAULT_INPUT, help="Path to a synthetic DOCX or PDF resume.")
    parser.add_argument("--mock", action="store_true", help="Use deterministic mock extraction instead of a real LLM.")
    parser.add_argument("--job-description", type=str, default=None, help="Optional job description text.")
    args = parser.parse_args()

    resume_path = args.input
    if resume_path == DEFAULT_INPUT and not resume_path.exists():
        create_synthetic_resume(resume_path)

    resume_text = read_resume_text(resume_path)
    llm_client = MockLLMClient() if args.mock else build_llm_client()
    profile = extract_candidate_profile(resume_text, llm_client, args.job_description)
    followup_message = generate_followup_message(profile)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    (OUTPUT_DIR / "raw_extracted_text.txt").write_text(resume_text, encoding="utf-8")
    (OUTPUT_DIR / "candidate_profile.json").write_text(
        json.dumps(profile.model_dump(mode="json"), indent=2),
        encoding="utf-8",
    )
    (OUTPUT_DIR / "missing_fields.json").write_text(
        json.dumps([field.model_dump() for field in profile.missing_fields], indent=2),
        encoding="utf-8",
    )
    (OUTPUT_DIR / "followup_email.txt").write_text(followup_message, encoding="utf-8")

    print(f"Processed: {resume_path}")
    print(f"Candidate: {profile.full_name or 'Unknown'}")
    print(f"Missing fields: {', '.join(field.label for field in profile.missing_fields) or 'None'}")
    print(f"Artifacts: {OUTPUT_DIR}")


def read_resume_text(path: Path) -> str:
    if path.suffix.lower() == ".pdf":
        return read_pdf_text(path)
    return read_docx(path).plain_text


def create_synthetic_resume(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading("Alex Morgan", level=1)
    document.add_paragraph("alex.morgan@example.com | +1 555 010 7722")
    document.add_paragraph("Location: Austin, TX")
    document.add_paragraph("https://www.linkedin.com/in/alex-morgan")
    document.add_heading("Summary", level=2)
    document.add_paragraph("Revenue operations analyst with experience improving recruiting funnels and CRM reporting.")
    document.add_paragraph("Skills: SQL, Python, Salesforce, Tableau")
    document.add_paragraph("Languages: English, Spanish")
    document.add_heading("Experience", level=2)
    table = document.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "Company"
    table.rows[0].cells[1].text = "Role"
    table.rows[0].cells[2].text = "Dates"
    table.rows[1].cells[0].text = "Northstar Talent"
    table.rows[1].cells[1].text = "Revenue Operations Analyst"
    table.rows[1].cells[2].text = "2022 - Present"
    document.save(path)


if __name__ == "__main__":
    main()
