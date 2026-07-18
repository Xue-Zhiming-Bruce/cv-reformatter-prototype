from pathlib import Path
import base64
import io

import pytest
from docx import Document
from docx.shared import Inches

from app.ingestion.docx_reader import read_docx
from app.ingestion.file_validator import CorruptedFileError, UnsupportedFileTypeError, validate_docx_file


def test_read_docx_extracts_paragraphs_and_tables(tmp_path: Path) -> None:
    path = tmp_path / "resume.docx"
    document = Document()
    document.add_paragraph("Jane Candidate")
    document.add_paragraph("Skills: Python, SQL")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Company"
    table.rows[0].cells[1].text = "Role"
    table.rows[1].cells[0].text = "Acme"
    table.rows[1].cells[1].text = "Engineer"
    document.save(path)

    content = read_docx(path)

    assert content.paragraphs == ["Jane Candidate", "Skills: Python, SQL"]
    assert content.tables[0].rows[1] == ["Acme", "Engineer"]
    assert "Acme | Engineer" in content.plain_text


def test_read_docx_with_embedded_image_preserves_text_extraction(tmp_path: Path) -> None:
    path = tmp_path / "resume-with-image.docx"
    document = Document()
    document.add_paragraph("Jane Candidate")
    document.add_picture(io.BytesIO(_one_pixel_png()), width=Inches(0.25))
    document.add_paragraph("Skills: Python, SQL")
    document.save(path)

    content = read_docx(path)

    assert content.paragraphs == ["Jane Candidate", "Skills: Python, SQL"]
    assert content.plain_text == "Jane Candidate\nSkills: Python, SQL"


def test_validate_docx_rejects_unsupported_extension(tmp_path: Path) -> None:
    path = tmp_path / "resume.txt"
    path.write_text("not a docx", encoding="utf-8")

    with pytest.raises(UnsupportedFileTypeError, match="Only .docx"):
        validate_docx_file(path)


def test_validate_docx_rejects_corrupted_file(tmp_path: Path) -> None:
    path = tmp_path / "resume.docx"
    path.write_text("not a zip", encoding="utf-8")

    with pytest.raises(CorruptedFileError, match="corrupted"):
        validate_docx_file(path)


def _one_pixel_png() -> bytes:
    return base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk+M/wHwAF/gL+X3mNvwAAAABJRU5ErkJggg=="
    )
