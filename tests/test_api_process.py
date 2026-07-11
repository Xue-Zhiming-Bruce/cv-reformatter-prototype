from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient
from pypdf import PdfWriter
from pytest import MonkeyPatch

from app import main
from app.main import app

client = TestClient(app)


@pytest.fixture(autouse=True)
def _use_tmp_local_database(tmp_path: Path, monkeypatch: MonkeyPatch) -> None:
    monkeypatch.setattr(main, "LOCAL_DATABASE_PATH", tmp_path / "local.sqlite3")


def _build_resume_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("Jane Candidate")
    document.add_paragraph("jane@example.com")
    document.add_paragraph("Location: Boston")
    document.add_paragraph("Skills: Python, SQL")
    document.save(path)


def _build_text_pdf(path: Path) -> None:
    lines = [
        "Jane Candidate",
        "jane@example.com",
        "Location: Boston",
        "Skills: Python, SQL",
    ]
    operations = ["BT /F1 12 Tf 72 720 Td"]
    for index, line in enumerate(lines):
        if index:
            operations.append("0 -16 Td")
        operations.append(f"({_escape_pdf_text(line)}) Tj")
    operations.append("ET")
    content = " ".join(operations).encode("ascii")

    objects = [
        b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n",
        b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n",
        b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >> endobj\n",
        b"4 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n",
        (
            b"5 0 obj << /Length "
            + str(len(content)).encode("ascii")
            + b" >> stream\n"
            + content
            + b"\nendstream endobj\n"
        ),
    ]
    pdf = b"%PDF-1.4\n"
    offsets = [0]
    for obj in objects:
        offsets.append(len(pdf))
        pdf += obj
    xref_start = len(pdf)
    pdf += f"xref\n0 {len(objects) + 1}\n".encode("ascii")
    pdf += b"0000000000 65535 f \n"
    for offset in offsets[1:]:
        pdf += f"{offset:010d} 00000 n \n".encode("ascii")
    pdf += (
        f"trailer << /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_start}\n%%EOF\n"
    ).encode("ascii")
    path.write_bytes(pdf)


def _escape_pdf_text(value: str) -> str:
    return value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def test_process_returns_profile_and_ledger(tmp_path: Path, monkeypatch: MonkeyPatch) -> None:
    monkeypatch.setenv("API_LLM_PROVIDER", "mock")
    monkeypatch.setattr(main, "GENERATED_OUTPUTS_DIR", tmp_path)
    monkeypatch.setattr(main, "export_docx_to_pdf", _fake_pdf_export)
    resume_path = tmp_path / "resume.docx"
    _build_resume_docx(resume_path)

    with resume_path.open("rb") as resume_file:
        response = client.post(
            "/api/process",
            files={
                "file": (
                    "resume.docx",
                    resume_file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert response.status_code == 200
    body = response.json()
    artifact_dir = tmp_path / body["artifact_id"]
    assert body["profile"]["full_name"] == "Jane Candidate"
    assert body["profile"]["email"] == "jane@example.com"
    assert body["ledger"]["needs_review"] == len(body["profile"]["missing_fields"])
    assert body["ledger"]["extracted"] == body["ledger"]["placed"]
    assert "Jane Candidate" in body["original_text"]
    assert body["original_pdf_preview_url"] == f"/api/artifacts/{body['artifact_id']}/original_resume_preview.pdf"
    assert body["original_preview_error"] is None
    assert body["artifact_metadata_url"] == f"/api/artifacts/{body['artifact_id']}/metadata"
    assert (artifact_dir / "raw_extracted_text.txt").read_text(encoding="utf-8") == body["original_text"]
    assert (artifact_dir / "candidate_profile.json").exists()
    assert (artifact_dir / "missing_fields.json").exists()
    assert (artifact_dir / "source_resume.docx").exists()
    assert (artifact_dir / "original_resume_preview.pdf").exists()

    preview_response = client.get(body["original_pdf_preview_url"])
    metadata_response = client.get(body["artifact_metadata_url"])
    assert preview_response.status_code == 200
    assert preview_response.headers["content-type"] == "application/pdf"
    assert preview_response.headers["content-disposition"].startswith("inline;")
    assert metadata_response.status_code == 200
    metadata = metadata_response.json()
    assert metadata["artifact_id"] == body["artifact_id"]
    assert metadata["status"] == "processed"
    assert metadata["source_filename"] == "resume.docx"
    assert metadata["source_file_type"] == "docx"
    assert metadata["original_pdf_preview_url"] == body["original_pdf_preview_url"]
    assert metadata["needs_review_count"] == len(body["profile"]["missing_fields"])
    assert metadata["debug_artifacts"]["candidate_profile"] == str(artifact_dir / "candidate_profile.json")


def test_process_accepts_text_pdf_upload(tmp_path: Path, monkeypatch: MonkeyPatch) -> None:
    monkeypatch.setenv("API_LLM_PROVIDER", "mock")
    monkeypatch.setattr(main, "GENERATED_OUTPUTS_DIR", tmp_path)
    resume_path = tmp_path / "resume.pdf"
    _build_text_pdf(resume_path)

    with resume_path.open("rb") as resume_file:
        response = client.post(
            "/api/process",
            files={"file": ("resume.pdf", resume_file, "application/pdf")},
        )

    assert response.status_code == 200
    body = response.json()
    artifact_dir = tmp_path / body["artifact_id"]
    assert body["profile"]["full_name"] == "Jane Candidate"
    assert body["profile"]["email"] == "jane@example.com"
    assert body["profile"]["skills"] == ["Python", "SQL"]
    assert "Jane Candidate" in body["original_text"]
    assert body["original_pdf_preview_url"] == f"/api/artifacts/{body['artifact_id']}/original_resume_preview.pdf"
    assert body["original_preview_error"] is None
    assert body["artifact_metadata_url"] == f"/api/artifacts/{body['artifact_id']}/metadata"
    assert (artifact_dir / "source_resume.pdf").exists()
    assert (artifact_dir / "original_resume_preview.pdf").read_bytes() == resume_path.read_bytes()

    metadata_response = client.get(body["artifact_metadata_url"])
    assert metadata_response.status_code == 200
    preview_response = client.get(body["original_pdf_preview_url"])
    assert preview_response.status_code == 200
    assert preview_response.headers["content-type"] == "application/pdf"
    assert preview_response.headers["content-disposition"].startswith("inline;")
    metadata = metadata_response.json()
    assert metadata["status"] == "processed"
    assert metadata["source_file_type"] == "pdf"


def test_process_rejects_pdf_without_extractable_text(tmp_path: Path) -> None:
    pdf_path = tmp_path / "blank.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    with pdf_path.open("wb") as output:
        writer.write(output)

    with pdf_path.open("rb") as resume_file:
        response = client.post(
            "/api/process",
            files={"file": ("blank.pdf", resume_file, "application/pdf")},
        )

    assert response.status_code == 400
    assert "no extractable text" in response.json()["detail"].lower()
    assert "scanned pdfs are not supported" in response.json()["detail"].lower()


def test_process_rejects_unsupported_upload(tmp_path: Path) -> None:
    text_path = tmp_path / "resume.txt"
    text_path.write_text("not a docx", encoding="utf-8")

    with text_path.open("rb") as text_file:
        response = client.post("/api/process", files={"file": ("resume.txt", text_file, "text/plain")})

    assert response.status_code == 400
    assert ".docx" in response.json()["detail"].lower()
    assert ".pdf" in response.json()["detail"].lower()


def _fake_pdf_export(docx_path: str | Path, output_dir: str | Path | None = None) -> Path:
    source_path = Path(docx_path)
    resolved_output_dir = Path(output_dir) if output_dir is not None else source_path.parent
    pdf_path = resolved_output_dir / f"{source_path.stem}.pdf"
    pdf_path.write_bytes(b"%PDF-1.4\n")
    return pdf_path
