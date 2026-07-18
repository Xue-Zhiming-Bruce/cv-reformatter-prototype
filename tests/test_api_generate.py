import json
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient
from pytest import MonkeyPatch

from app import main
from app.extraction.candidate_schema import CandidateProfile, WorkExperience
from app.generation.target_format_blueprint import get_target_format_blueprint
from app.main import app

client = TestClient(app)


@pytest.fixture(autouse=True)
def _use_tmp_local_database(tmp_path: Path, monkeypatch: MonkeyPatch) -> None:
    monkeypatch.setattr(main, "LOCAL_DATABASE_PATH", tmp_path / "local.sqlite3")


def test_generate_outputs_saves_artifacts_and_returns_downloads(tmp_path: Path, monkeypatch: MonkeyPatch) -> None:
    monkeypatch.setattr(main, "GENERATED_OUTPUTS_DIR", tmp_path)
    monkeypatch.setattr(main, "export_docx_to_pdf", _fake_pdf_export)
    profile = CandidateProfile(
        full_name="Jane Candidate",
        email="jane@example.com",
        location="Boston",
        professional_summary="Backend engineer focused on data platforms.",
        skills=["Python", "SQL"],
        work_experience=[
            WorkExperience(
                company="DataWorks",
                title="Senior Engineer",
                start_date="2020",
                end_date="2024",
                description=["Built analytics pipelines."],
            )
        ],
    )

    response = client.post(
        "/api/generate",
        json={
            "profile": profile.model_dump(mode="json"),
            "client_display_rules": {
                "salary_expectation": "available_upon_request",
                "notice_period": "pending_confirmation",
            },
            "blind_profile": True,
            "template_name": "apex_standard",
            "original_text": "Jane Candidate\nSkills: Python, SQL",
        },
    )

    assert response.status_code == 200
    body = response.json()
    artifact_dir = tmp_path / body["artifact_id"]
    assert (artifact_dir / "raw_extracted_text.txt").read_text(encoding="utf-8") == "Jane Candidate\nSkills: Python, SQL"
    assert not (artifact_dir / "candidate_profile.json").exists()
    assert not (artifact_dir / "missing_fields.json").exists()
    assert (artifact_dir / "export_snapshot.json").exists()
    assert (artifact_dir / "export_missing_fields.json").exists()
    assert (artifact_dir / "client_render_context.json").exists()
    assert (artifact_dir / "candidate_profile.docx").exists()
    assert (artifact_dir / "candidate_profile.pdf").exists()
    assert body["docx_download_url"] == f"/api/artifacts/{body['artifact_id']}/candidate_profile.docx?download=1"
    assert body["pdf_download_url"] == f"/api/artifacts/{body['artifact_id']}/candidate_profile.pdf?download=1"
    assert body["pdf_preview_url"] == f"/api/artifacts/{body['artifact_id']}/candidate_profile.pdf"
    assert body["artifact_metadata_url"] == f"/api/artifacts/{body['artifact_id']}/metadata"
    assert "Salary expectation" in body["followup_message"]
    assert {field["field_name"] for field in body["missing_fields"]} >= {
        "salary_expectation",
        "notice_period",
        "work_authorization",
        "interview_availability",
    }

    download_response = client.get(body["docx_download_url"])
    pdf_response = client.get(body["pdf_download_url"])
    metadata_response = client.get(body["artifact_metadata_url"])
    assert download_response.status_code == 200
    assert download_response.headers["content-disposition"].startswith("attachment;")
    assert pdf_response.status_code == 200
    assert pdf_response.headers["content-type"] == "application/pdf"
    assert pdf_response.headers["content-disposition"].startswith("attachment;")
    preview_response = client.get(body["pdf_preview_url"])
    assert preview_response.headers["content-disposition"].startswith("inline;")
    assert metadata_response.status_code == 200
    metadata = metadata_response.json()
    assert metadata["status"] == "generated"
    assert metadata["template_name"] == "apex_standard"
    assert metadata["blind_profile"] is True
    assert metadata["docx_download_url"] == body["docx_download_url"]
    assert metadata["pdf_preview_url"] == body["pdf_preview_url"]
    assert metadata["needs_review_count"] == len(body["missing_fields"])


def test_process_then_generate_preserves_original_profile_and_exports_timeline_edits(
    tmp_path: Path,
    monkeypatch: MonkeyPatch,
) -> None:
    monkeypatch.setenv("API_LLM_PROVIDER", "mock")
    monkeypatch.setattr(main, "GENERATED_OUTPUTS_DIR", tmp_path)
    monkeypatch.setattr(main, "export_docx_to_pdf", _fake_pdf_export)
    resume_path = tmp_path / "timeline-resume.docx"
    resume = Document()
    resume.add_paragraph("Jane Candidate")
    resume.add_paragraph("jane@example.com")
    resume.add_paragraph("Location: Boston")
    resume.add_paragraph("Summary")
    resume.add_paragraph("Backend engineer focused on reliable systems.")
    resume.add_paragraph("Skills")
    resume.add_paragraph("Python, SQL")
    resume.add_paragraph("Experience")
    resume.add_paragraph("Senior Engineer at DataWorks (2020 - 2024)")
    resume.add_paragraph("Built analytics pipelines.")
    resume.add_paragraph("Education")
    resume.add_paragraph("BSc Computer Science - Example University (2019)")
    resume.save(resume_path)

    with resume_path.open("rb") as resume_file:
        process_response = client.post(
            "/api/process",
            files={
                "file": (
                    resume_path.name,
                    resume_file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert process_response.status_code == 200
    process_body = process_response.json()
    artifact_id = process_body["artifact_id"]
    artifact_dir = tmp_path / artifact_id
    original_profile_bytes = (artifact_dir / "candidate_profile.json").read_bytes()
    original_missing_fields_bytes = (artifact_dir / "missing_fields.json").read_bytes()
    original_text_bytes = (artifact_dir / "raw_extracted_text.txt").read_bytes()

    target_response = client.post(
        "/api/target-format",
        data={"artifact_id": artifact_id},
        files={"file": ("target-reference.pdf", b"%PDF-1.4\n", "application/pdf")},
    )
    assert target_response.status_code == 200
    target_format = target_response.json()["target_format"]

    edited_profile = process_body["profile"]
    edited_profile["work_experience"][0]["start_date"] = "January 2021"
    edited_profile["work_experience"][0]["end_date"] = "Present"
    edited_profile["work_experience"][0]["title"] = "Principal Engineer"

    generate_response = client.post(
        "/api/generate",
        json={
            "artifact_id": artifact_id,
            "profile": edited_profile,
            "template_name": "client_10554236_v1",
            "original_text": "this must not replace the preserved extraction",
            "target_format": target_format,
        },
    )

    assert generate_response.status_code == 200
    assert (artifact_dir / "candidate_profile.json").read_bytes() == original_profile_bytes
    assert (artifact_dir / "missing_fields.json").read_bytes() == original_missing_fields_bytes
    assert (artifact_dir / "raw_extracted_text.txt").read_bytes() == original_text_bytes

    export_snapshot = json.loads((artifact_dir / "export_snapshot.json").read_text(encoding="utf-8"))
    exported_experience = export_snapshot["work_experience"][0]
    assert exported_experience["start_date"] == "January 2021"
    assert exported_experience["end_date"] == "Present"
    assert exported_experience["title"] == "Principal Engineer"
    assert (artifact_dir / "export_missing_fields.json").exists()

    generated = Document(artifact_dir / "candidate_profile.docx")
    generated_text = "\n".join(paragraph.text for paragraph in generated.paragraphs)
    assert "Principal Engineer" in generated_text
    assert "January 2021 to Present" in generated_text
    assert (artifact_dir / "candidate_profile.pdf").exists()

    docx_download = client.get(generate_response.json()["docx_download_url"])
    pdf_download = client.get(generate_response.json()["pdf_download_url"])
    assert docx_download.status_code == 200
    assert docx_download.headers["content-disposition"].startswith("attachment;")
    assert pdf_download.status_code == 200
    assert pdf_download.headers["content-disposition"].startswith("attachment;")

    metadata = client.get(generate_response.json()["artifact_metadata_url"]).json()
    assert metadata["debug_artifacts"]["candidate_profile"] == str(
        artifact_dir / "candidate_profile.json"
    )
    assert metadata["debug_artifacts"]["export_snapshot"] == str(artifact_dir / "export_snapshot.json")


def test_download_generated_artifact_rejects_unknown_or_unsafe_paths(tmp_path: Path, monkeypatch: MonkeyPatch) -> None:
    monkeypatch.setattr(main, "GENERATED_OUTPUTS_DIR", tmp_path)

    unsafe_response = client.get("/api/artifacts/bad.id/candidate_profile.docx")
    unknown_file_response = client.get("/api/artifacts/generation_123/candidate_profile.json")

    assert unsafe_response.status_code == 404
    assert unknown_file_response.status_code == 404


def test_generate_outputs_can_reuse_process_artifact_with_target_format(
    tmp_path: Path,
    monkeypatch: MonkeyPatch,
) -> None:
    monkeypatch.setattr(main, "GENERATED_OUTPUTS_DIR", tmp_path)
    monkeypatch.setattr(main, "export_docx_to_pdf", _fake_pdf_export)
    artifact_id = "artifact_existing_session"
    artifact_dir = tmp_path / artifact_id
    artifact_dir.mkdir()
    profile = CandidateProfile(
        full_name="Jane Candidate",
        email="jane@example.com",
        location="Boston",
        skills=["Python", "SQL"],
    )

    target_response = client.post(
        "/api/target-format",
        data={"artifact_id": artifact_id},
        files={
            "file": (
                "client-format.pdf",
                b"%PDF-1.4\n",
                "application/pdf",
            )
        },
    )

    assert target_response.status_code == 200
    target_body = target_response.json()
    target_format = target_body["target_format"]
    assert target_body["artifact_metadata_url"] == f"/api/artifacts/{artifact_id}/metadata"

    response = client.post(
        "/api/generate",
        json={
            "artifact_id": artifact_id,
            "profile": profile.model_dump(mode="json"),
            "target_format": target_format,
            "original_text": "Jane Candidate\nSkills: Python, SQL",
        },
    )

    assert response.status_code == 200
    body = response.json()
    assert body["artifact_id"] == artifact_id
    assert body["pdf_preview_url"] == f"/api/artifacts/{artifact_id}/candidate_profile.pdf"
    assert (artifact_dir / "target_format_reference.pdf").exists()
    assert (artifact_dir / "target_format.json").exists()
    assert (artifact_dir / "candidate_profile.docx").exists()
    assert (artifact_dir / "candidate_profile.pdf").exists()

    metadata_response = client.get(body["artifact_metadata_url"])
    assert metadata_response.status_code == 200
    metadata = metadata_response.json()
    assert metadata["status"] == "generated"
    assert metadata["target_format_role"] == "pdf_reference"
    assert metadata["target_format_filename"] == "client-format.pdf"
    assert metadata["pdf_download_url"] == body["pdf_download_url"]


def test_target_format_upload_accepts_docx_templates_and_pdf_references(
    tmp_path: Path,
    monkeypatch: MonkeyPatch,
) -> None:
    monkeypatch.setattr(main, "GENERATED_OUTPUTS_DIR", tmp_path)
    artifact_id = "artifact_target_formats"
    docx_path = tmp_path / "target-template.docx"
    docx_path.write_bytes(b"synthetic docx bytes")

    with docx_path.open("rb") as template_file:
        docx_response = client.post(
            "/api/target-format",
            data={"artifact_id": artifact_id},
            files={
                "file": (
                    "target-template.docx",
                    template_file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    pdf_response = client.post(
        "/api/target-format",
        data={"artifact_id": artifact_id},
        files={"file": ("target-reference.pdf", b"%PDF-1.4\n", "application/pdf")},
    )

    assert docx_response.status_code == 200
    docx_body = docx_response.json()
    assert docx_body["artifact_id"] == artifact_id
    assert docx_body["artifact_metadata_url"] == f"/api/artifacts/{artifact_id}/metadata"
    assert docx_body["target_format"]["role"] == "docx_template"
    assert docx_body["target_format"]["download_url"] == f"/api/artifacts/{artifact_id}/target_format_template.docx"
    assert (tmp_path / artifact_id / "target_format_template.docx").exists()

    assert pdf_response.status_code == 200
    pdf_body = pdf_response.json()
    assert pdf_body["target_format"]["role"] == "pdf_reference"
    assert "does not reverse-engineer" in pdf_body["target_format"]["note"]
    assert (tmp_path / artifact_id / "target_format_reference.pdf").exists()

    download_response = client.get(pdf_body["target_format"]["download_url"])
    metadata_response = client.get(pdf_body["artifact_metadata_url"])
    assert download_response.status_code == 200
    assert metadata_response.status_code == 200
    metadata = metadata_response.json()
    assert metadata["status"] == "target_format_uploaded"
    assert metadata["target_format_role"] == "pdf_reference"
    assert metadata["target_format_filename"] == "target-reference.pdf"


def test_registered_pdf_reference_selects_manual_blueprint(
    tmp_path: Path,
    monkeypatch: MonkeyPatch,
) -> None:
    monkeypatch.setattr(main, "GENERATED_OUTPUTS_DIR", tmp_path)
    monkeypatch.setattr(main, "export_docx_to_pdf", _fake_pdf_export)
    blueprint = get_target_format_blueprint("client_10554236_v1")
    monkeypatch.setattr(main, "match_target_format_blueprint", lambda *_args, **_kwargs: blueprint)
    artifact_id = "artifact_registered_blueprint"
    target_response = client.post(
        "/api/target-format",
        data={"artifact_id": artifact_id},
        files={"file": ("known-reference.pdf", b"%PDF-1.4\n", "application/pdf")},
    )

    assert target_response.status_code == 200
    target_format = target_response.json()["target_format"]
    assert target_format["matched_template_id"] == "client_10554236_v1"
    assert target_format["blueprint_version"] == 1
    assert target_format["format_support"] == "manual_blueprint"
    assert target_format["generation_strategy"] == "registered_target_format_blueprint"

    profile = CandidateProfile(
        full_name="Jane Candidate",
        professional_summary="Accountant focused on reporting.",
        skills=["Reporting", "Reconciliations"],
    )
    generate_response = client.post(
        "/api/generate",
        json={
            "artifact_id": artifact_id,
            "profile": profile.model_dump(mode="json"),
            "target_format": target_format,
        },
    )

    assert generate_response.status_code == 200
    assert generate_response.json()["template_id"] == "client_10554236_v1"
    generated = Document(tmp_path / artifact_id / "candidate_profile.docx")
    assert generated.styles["Normal"].font.name == "Times New Roman"


def test_artifact_metadata_list_returns_recent_jobs(tmp_path: Path, monkeypatch: MonkeyPatch) -> None:
    monkeypatch.setattr(main, "GENERATED_OUTPUTS_DIR", tmp_path)

    target_response = client.post(
        "/api/target-format",
        data={"artifact_id": "artifact_listed"},
        files={"file": ("target-reference.pdf", b"%PDF-1.4\n", "application/pdf")},
    )

    assert target_response.status_code == 200
    response = client.get("/api/artifacts")

    assert response.status_code == 200
    artifacts = response.json()["artifacts"]
    assert len(artifacts) == 1
    assert artifacts[0]["artifact_id"] == "artifact_listed"
    assert artifacts[0]["status"] == "target_format_uploaded"


def test_target_format_upload_rejects_unsupported_files(tmp_path: Path, monkeypatch: MonkeyPatch) -> None:
    monkeypatch.setattr(main, "GENERATED_OUTPUTS_DIR", tmp_path)

    response = client.post(
        "/api/target-format",
        files={"file": ("target.txt", b"not a supported template", "text/plain")},
    )

    assert response.status_code == 400
    assert ".docx" in response.json()["detail"].lower()
    assert ".pdf" in response.json()["detail"].lower()


def test_generate_rejects_missing_target_format_reference(tmp_path: Path, monkeypatch: MonkeyPatch) -> None:
    monkeypatch.setattr(main, "GENERATED_OUTPUTS_DIR", tmp_path)
    profile = CandidateProfile(full_name="Jane Candidate")
    artifact_id = "artifact_missing_target"

    response = client.post(
        "/api/generate",
        json={
            "artifact_id": artifact_id,
            "profile": profile.model_dump(mode="json"),
            "target_format": {
                "artifact_id": artifact_id,
                "input_filename": "target-reference.pdf",
                "stored_filename": "target_format_reference.pdf",
                "input_type": "pdf",
                "role": "pdf_reference",
                "generation_strategy": "controlled_docx_renderer",
                "accepted_for_generation": True,
                "used_as_template_source": False,
                "download_url": f"/api/artifacts/{artifact_id}/target_format_reference.pdf",
                "note": "PDF target format stored as a visual reference only.",
            },
        },
    )

    assert response.status_code == 400
    assert "has not been uploaded" in response.json()["detail"]


def _fake_pdf_export(docx_path: str | Path, output_dir: str | Path | None = None) -> Path:
    source_path = Path(docx_path)
    resolved_output_dir = Path(output_dir) if output_dir is not None else source_path.parent
    pdf_path = resolved_output_dir / f"{source_path.stem}.pdf"
    pdf_path.write_bytes(b"%PDF-1.4\n")
    return pdf_path
