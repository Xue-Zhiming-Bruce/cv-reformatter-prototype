from pathlib import Path

import pytest
from docx import Document

from app.extraction.candidate_schema import CandidateProfile, DisplayRule, WorkExperience
from app.generation.docx_renderer import render_docx
from app.generation.target_format_blueprint import get_target_format_blueprint
from app.generation.template_mapper import build_client_render_context


def test_docx_renderer_writes_client_facing_profile(tmp_path: Path) -> None:
    profile = CandidateProfile(
        full_name="Jane Candidate",
        location="Boston",
        professional_summary="Backend engineer focused on data platforms.",
        skills=["Python", "SQL"],
        work_experience=[
            WorkExperience(
                company="DataWorks",
                title="Senior Engineer",
                start_date="2020",
                end_date="2024",
                description=["Built analytics pipelines."],
            )
        ],
        salary_expectation=None,
        client_display_rules={"salary_expectation": DisplayRule.AVAILABLE_UPON_REQUEST},
    )
    context = build_client_render_context(profile)

    output_path = render_docx(context, tmp_path / "candidate_profile.docx")

    assert output_path.exists()
    text = _extract_docx_text(output_path)
    assert "Jane Candidate" in text
    assert "Backend engineer focused on data platforms." in text
    assert "Python" in text
    assert "Senior Engineer | DataWorks | 2020 - 2024" in text
    assert "Salary expectation: Available upon request" in text
    assert "Missing" not in text


def test_docx_renderer_accepts_mapping_context(tmp_path: Path) -> None:
    context = {
        "candidate_heading": "Candidate A",
        "candidate_subheading": "Backend Lead",
        "contact_lines": ["Location: To be confirmed"],
        "additional_details": [{"label": "Notice period", "value": "To be confirmed"}],
    }

    output_path = render_docx(context, tmp_path / "candidate_profile.docx")

    assert "Candidate A" in _extract_docx_text(output_path)


def test_classic_blueprint_controls_layout_typography_and_section_order(tmp_path: Path) -> None:
    profile = CandidateProfile(
        full_name="Jane Candidate",
        professional_summary="Financial accountant focused on reporting and analysis.",
        skills=["Account reconciliations", "Financial reporting", "ERP"],
        work_experience=[
            WorkExperience(
                company="Company Name",
                title="Accountant",
                location="City, State",
                start_date="July 2011",
                end_date="November 2012",
                description=["Resolved general ledger posting differences."],
            )
        ],
    )
    blueprint = get_target_format_blueprint("client_10554236_v1")
    context = build_client_render_context(profile, template_name=blueprint.template_id)

    output_path = render_docx(context, tmp_path / "classic.docx", blueprint=blueprint)

    document = Document(output_path)
    section = document.sections[0]
    text = _extract_docx_text(output_path)
    assert section.page_width.inches == pytest.approx(blueprint.page.width_inches, abs=0.01)
    assert section.left_margin.inches == pytest.approx(blueprint.page.left_margin_inches, abs=0.01)
    assert document.styles["Normal"].font.name == "Times New Roman"
    assert document.styles["Normal"].font.size.pt == pytest.approx(10.0)
    assert text.index("Summary") < text.index("Experience")
    assert "Company Name July 2011 to November 2012 Accountant" in text
    assert "City, State" in text
    assert len(document.tables) == 1
    assert {style.color_hex for style in blueprint.typography.values()} == {"000000"}

    paragraphs = list(document.paragraphs)
    paragraphs.extend(
        paragraph
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    )
    direct_run_colors = {
        str(run.font.color.rgb)
        for paragraph in paragraphs
        for run in paragraph.runs
        if run.text and run.font.color.rgb is not None
    }
    assert direct_run_colors == {"000000"}


def _extract_docx_text(path: Path) -> str:
    document = Document(path)
    return "\n".join(paragraph.text for paragraph in document.paragraphs)
